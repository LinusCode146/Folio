from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

doc = Document()
doc.add_heading('Q1/2 Wirtschaftspolitische Konzeptionen im Vergleich', level=1)

rows_data = [
    ("Konzept",
     "Nachfrageorientierte Wirtschaftspolitik\nNach J. M. Keynes",
     "Angebotsorientierte Wirtschaftspolitik\nNach M. Friedman"),
    ("Krisenerklärung und Analyse",
     "Krisen entstehen durch zu geringe gesamtwirtschaftliche Nachfrage. Der Markt reguliert sich nicht von allein und bleibt im Unterbeschäftigungsgleichgewicht stecken.",
     "Krisen entstehen durch zu viele Staatseingriffe, zu hohe Steuern/Löhne und eine falsche Geldpolitik. Zu starke Geldmengenausweitung führt zu Inflation."),
    ("Entstehungszusammenhang / Historie",
     "Reaktion auf die Weltwirtschaftskrise 1929; Hauptwerk 1936. In Deutschland umgesetzt durch K. Schiller mit dem Stabilitätsgesetz 1967.",
     "Entstanden in den 1970ern als Antwort auf die Stagflation. Umgesetzt in den 1980ern unter Reagan („Reaganomics“) und Thatcher."),
    ("Wirtschaftspolitische Ziele",
     "Vollbeschäftigung und stetiges Wachstum durch antizyklische Konjunktursteuerung. Orientierung am „magischen Viereck“.",
     "Preisniveaustabilität und Stärkung der Wettbewerbsfähigkeit. Wachstum und Beschäftigung folgen laut Say’schem Theorem von selbst."),
    ("Möglichkeiten der Umsetzung",
     "Aktive Fiskalpolitik: in der Rezession deficit spending, Steuersenkungen, expansive Geldpolitik; im Boom Ausgaben senken und Steuern erhöhen.",
     "Staat zurücknehmen: Steuern senken, Subventionen abbauen, Deregulierung, Privatisierung, flexibler Arbeitsmarkt, feste Geldmengenregel."),
    ("Beabsichtigte Wirkungsweise",
     "Höhere Staatsnachfrage steigert über den Multiplikatoreffekt Einkommen und Konsum. Unternehmen investieren, stellen ein – die Wirtschaft kommt aus der Krise.",
     "Niedrige Steuern erhöhen Gewinne und Investitionen. Mehr Produktion schafft Arbeitsplätze; Wohlstand „sickert“ nach unten (trickle-down)."),
    ("Probleme und Kritik",
     "Hohe Staatsverschuldung, Time-lags und prozyklische Wirkung, Inflationsgefahr, Crowding-out-Effekt; versagt bei Angebotsschocks (Stagflation).",
     "Wachsende soziale Ungleichheit, trickle-down kaum belegbar, Sozialabbau schwächt Nachfrage; versagt in echten Nachfragekrisen (z. B. 2008)."),
]

table = doc.add_table(rows=len(rows_data), cols=3)
table.style = 'Table Grid'

for i, (cat, keynes, friedman) in enumerate(rows_data):
    row = table.rows[i]
    row.cells[0].text = cat
    row.cells[1].text = keynes
    row.cells[2].text = friedman
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True

for i, row in enumerate(table.rows):
    if i == 0:
        continue
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True

out = '/Users/linusschulze/RustRoverProjects/writable/Wirtschaftspolitische_Konzeptionen_Loesung.docx'
doc.save(out)
print(out)
